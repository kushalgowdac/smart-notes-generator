"""
Export Notes to Multiple Formats
Convert generated lecture notes to PDF, DOCX, Anki flashcards, or Notion-ready format

Usage:
    python export_notes.py data/lectures/my_lecture --format pdf
    python export_notes.py data/lectures/my_lecture --format pdf,docx,anki
    python export_notes.py data/lectures/my_lecture --format all
"""

import os
import json
import re
from pathlib import Path
import argparse
import logging
from datetime import datetime

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class NotesExporter:
    """Export notes to various formats"""
    
    def __init__(self, lecture_dir):
        self.lecture_path = Path(lecture_dir)
        
        if not self.lecture_path.exists():
            raise FileNotFoundError(f"Lecture directory not found: {lecture_dir}")
        
        self.notes_file = self.lecture_path / 'notes.md'
        
        if not self.notes_file.exists():
            raise FileNotFoundError(f"notes.md not found in {lecture_dir}")
        
        # Load notes
        with open(self.notes_file, 'r', encoding='utf-8') as f:
            self.markdown_content = f.read()
        
        self.lecture_name = self.lecture_path.name
        self.output_dir = self.lecture_path / 'exports'
        self.output_dir.mkdir(exist_ok=True)
    
    def export_pdf(self):
        """Export to PDF using markdown2pdf or weasyprint"""
        logger.info("Exporting to PDF...")
        
        try:
            import markdown
            from weasyprint import HTML, CSS
            
            # Convert markdown to HTML
            html_content = markdown.markdown(
                self.markdown_content,
                extensions=['extra', 'codehilite', 'toc', 'tables', 'fenced_code']
            )
            
            # CSS styling
            css = CSS(string="""
                @page {
                    margin: 2.5cm;
                    @top-right {
                        content: counter(page);
                        font-size: 10pt;
                        color: #666;
                    }
                }
                body {
                    font-family: 'Segoe UI', 'Arial', sans-serif;
                    line-height: 1.8;
                    color: #2c3e50;
                    font-size: 11pt;
                }
                h1 {
                    color: #2980b9;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 12px;
                    margin-top: 30px;
                    font-size: 24pt;
                }
                h2 {
                    color: #27ae60;
                    margin-top: 25px;
                    font-size: 18pt;
                    border-left: 4px solid #27ae60;
                    padding-left: 12px;
                }
                h3 {
                    color: #8e44ad;
                    font-size: 14pt;
                    margin-top: 20px;
                }
                img {
                    max-width: 90%;
                    border: 1px solid #bdc3c7;
                    padding: 8px;
                    margin: 15px auto;
                    display: block;
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                code {
                    background: #ecf0f1;
                    padding: 3px 7px;
                    border-radius: 4px;
                    font-family: 'Consolas', 'Monaco', monospace;
                    font-size: 10pt;
                    color: #e74c3c;
                }
                pre {
                    background: #2c3e50;
                    color: #ecf0f1;
                    padding: 15px;
                    border-radius: 6px;
                    overflow-x: auto;
                    font-size: 9pt;
                    line-height: 1.6;
                }
                pre code {
                    background: transparent;
                    color: #ecf0f1;
                    padding: 0;
                }
                blockquote {
                    border-left: 4px solid #f39c12;
                    padding-left: 20px;
                    background: #fef5e7;
                    margin: 15px 0;
                    padding: 12px 20px;
                    font-style: italic;
                    color: #7d6608;
                }
                table {
                    border-collapse: collapse;
                    width: 100%;
                    margin: 20px 0;
                }
                th {
                    background: #3498db;
                    color: white;
                    padding: 12px;
                    text-align: left;
                    font-weight: bold;
                }
                td {
                    border: 1px solid #bdc3c7;
                    padding: 10px;
                }
                tr:nth-child(even) {
                    background: #ecf0f1;
                }
                ul, ol {
                    margin: 12px 0;
                    padding-left: 30px;
                }
                li {
                    margin: 8px 0;
                }
                a {
                    color: #3498db;
                    text-decoration: none;
                }
                a:hover {
                    text-decoration: underline;
                }
                hr {
                    border: none;
                    border-top: 2px solid #ecf0f1;
                    margin: 30px 0;
                }
            """)
            
            # Add title page
            html_with_cover = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{self.lecture_name} - Lecture Notes</title>
            </head>
            <body>
                <div style="text-align: center; margin-top: 200px;">
                    <h1 style="font-size: 36pt; color: #2980b9; border: none;">{self.lecture_name.replace('_', ' ').title()}</h1>
                    <h2 style="font-size: 18pt; color: #7f8c8d; border: none; padding: 0;">Lecture Notes</h2>
                    <p style="font-size: 12pt; color: #95a5a6; margin-top: 50px;">
                        Generated: {datetime.now().strftime('%B %d, %Y')}
                    </p>
                </div>
                <div style="page-break-after: always;"></div>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF
            pdf_path = self.output_dir / f'{self.lecture_name}_notes.pdf'
            HTML(string=html_with_cover, base_url=str(self.lecture_path)).write_pdf(
                pdf_path,
                stylesheets=[css]
            )
            
            logger.info(f"✓ PDF saved: {pdf_path}")
            return pdf_path
        
        except ImportError as e:
            logger.error("PDF export failed - missing dependencies")
            logger.info("Install with: pip install markdown weasyprint")
            logger.info("Note: On Windows, you may also need GTK: https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases")
            return None
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            return None
    
    def export_docx(self):
        """Export to Microsoft Word DOCX"""
        logger.info("Exporting to DOCX...")
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title page
            title = doc.add_heading(self.lecture_name.replace('_', ' ').title(), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph('Lecture Notes')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(18)
            subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(12)
            
            doc.add_page_break()
            
            # Parse markdown and add content
            lines = self.markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    continue
                
                # Headers
                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                
                # Lists
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\. ', line):
                    doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
                
                # Code blocks
                elif line.startswith('```'):
                    continue  # Skip code block markers
                
                # Images
                elif line.startswith('!['):
                    # Extract image path
                    match = re.search(r'!\[.*?\]\((.*?)\)', line)
                    if match:
                        img_path = self.lecture_path / match.group(1)
                        if img_path.exists():
                            doc.add_picture(str(img_path), width=Inches(6))
                
                # Horizontal rule
                elif line == '---':
                    doc.add_paragraph('_' * 50)
                
                # Regular paragraph
                else:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.line_spacing = 1.5
            
            # Save
            docx_path = self.output_dir / f'{self.lecture_name}_notes.docx'
            doc.save(docx_path)
            
            logger.info(f"✓ DOCX saved: {docx_path}")
            return docx_path
        
        except ImportError:
            logger.error("DOCX export failed - python-docx not installed")
            logger.info("Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            return None
    
    def export_anki(self):
        """Export as Anki flashcards (APKG file)"""
        logger.info("Exporting to Anki flashcards...")
        
        try:
            import genanki
            
            # Create deck
            deck_id = hash(self.lecture_name) % (10 ** 10)
            deck = genanki.Deck(deck_id, f"{self.lecture_name.replace('_', ' ').title()} - Lecture Notes")
            
            # Simple model (front/back)
            model = genanki.Model(
                1607392319,
                'Lecture Notes Model',
                fields=[
                    {'name': 'Question'},
                    {'name': 'Answer'},
                    {'name': 'Slide'},
                ],
                templates=[
                    {
                        'name': 'Card 1',
                        'qfmt': '<div style="font-size: 20px;">{{Question}}</div><br><small>Slide: {{Slide}}</small>',
                        'afmt': '{{FrontSide}}<hr><div style="font-size: 18px;">{{Answer}}</div>',
                    },
                ])
            
            # Extract Q&A from notes
            # Strategy: Use headings as questions, content as answers
            lines = self.markdown_content.split('\n')
            
            current_slide = "Unknown"
            current_question = None
            current_answer = []
            
            for line in lines:
                line = line.strip()
                
                # Track slide number
                if line.startswith('## Slide '):
                    current_slide = line[3:]
                
                # Level 3/4 headings = questions
                elif line.startswith('### ') or line.startswith('#### '):
                    # Save previous card
                    if current_question and current_answer:
                        answer_text = '\n'.join(current_answer).strip()
                        if answer_text:
                            note = genanki.Note(
                                model=model,
                                fields=[current_question, answer_text, current_slide]
                            )
                            deck.add_note(note)
                    
                    # New question
                    current_question = line.replace('###', '').replace('####', '').strip()
                    current_answer = []
                
                # Accumulate answer
                elif line and not line.startswith('#') and not line.startswith('---'):
                    current_answer.append(line)
            
            # Save last card
            if current_question and current_answer:
                answer_text = '\n'.join(current_answer).strip()
                if answer_text:
                    note = genanki.Note(
                        model=model,
                        fields=[current_question, answer_text, current_slide]
                    )
                    deck.add_note(note)
            
            # Generate package
            anki_path = self.output_dir / f'{self.lecture_name}_flashcards.apkg'
            genanki.Package(deck).write_to_file(anki_path)
            
            logger.info(f"✓ Anki deck saved: {anki_path}")
            logger.info(f"  Total flashcards: {len(deck.notes)}")
            return anki_path
        
        except ImportError:
            logger.error("Anki export failed - genanki not installed")
            logger.info("Install with: pip install genanki")
            return None
        except Exception as e:
            logger.error(f"Anki export error: {e}")
            return None
    
    def export_notion(self):
        """Export as Notion-ready Markdown"""
        logger.info("Exporting for Notion...")
        
        # Notion uses standard markdown, but we need to clean up image paths
        # to be web-accessible or uploaded separately
        
        notion_md = self.markdown_content
        
        # Add Notion-specific metadata header
        header = f"""---
title: {self.lecture_name.replace('_', ' ').title()}
created: {datetime.now().isoformat()}
tags: lecture-notes, {self.lecture_name}
---

"""
        
        notion_md = header + notion_md
        
        # Add note about images
        image_note = """
> **Note about images:** 
> Images are referenced locally. When importing to Notion:
> 1. Upload slides to Notion
> 2. Replace image links manually, or
> 3. Use Notion's "Upload" feature to import with images
---

"""
        
        notion_md = notion_md.replace('# ', image_note + '\n# ', 1)
        
        # Save
        notion_path = self.output_dir / f'{self.lecture_name}_notion.md'
        with open(notion_path, 'w', encoding='utf-8') as f:
            f.write(notion_md)
        
        logger.info(f"✓ Notion markdown saved: {notion_path}")
        logger.info("  Import instructions:")
        logger.info("    1. Open Notion → Import → Markdown")
        logger.info(f"    2. Select {notion_path.name}")
        logger.info("    3. Upload slide images separately if needed")
        
        return notion_path


def main():
    parser = argparse.ArgumentParser(
        description='Export lecture notes to multiple formats',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python export_notes.py data/lectures/deadlock_os --format pdf
  python export_notes.py data/lectures/chemistry_01 --format pdf,docx
  python export_notes.py data/lectures/algo_1 --format all

Supported Formats:
  pdf    - Professional PDF with styling
  docx   - Microsoft Word document
  anki   - Anki flashcard deck (.apkg)
  notion - Notion-ready Markdown
  all    - Export to all formats

Dependencies:
  pdf    → pip install markdown weasyprint
  docx   → pip install python-docx
  anki   → pip install genanki
  notion → (no extra dependencies)
        """
    )
    
    parser.add_argument('lecture_dir', help='Path to lecture directory with notes.md')
    parser.add_argument('--format', '-f', default='pdf', 
                       help='Export format(s): pdf, docx, anki, notion, or all (comma-separated)')
    
    args = parser.parse_args()
    
    try:
        exporter = NotesExporter(args.lecture_dir)
        
        # Parse formats
        formats = args.format.lower().split(',')
        if 'all' in formats:
            formats = ['pdf', 'docx', 'anki', 'notion']
        
        logger.info("="*70)
        logger.info(f"EXPORTING NOTES: {exporter.lecture_name}")
        logger.info("="*70)
        logger.info(f"Formats: {', '.join(formats)}")
        logger.info("")
        
        results = {}
        
        if 'pdf' in formats:
            results['pdf'] = exporter.export_pdf()
        
        if 'docx' in formats:
            results['docx'] = exporter.export_docx()
        
        if 'anki' in formats:
            results['anki'] = exporter.export_anki()
        
        if 'notion' in formats:
            results['notion'] = exporter.export_notion()
        
        # Summary
        logger.info("")
        logger.info("="*70)
        logger.info("EXPORT SUMMARY")
        logger.info("="*70)
        
        success_count = sum(1 for path in results.values() if path is not None)
        
        for fmt, path in results.items():
            if path:
                logger.info(f"✓ {fmt.upper()}: {path}")
            else:
                logger.info(f"✗ {fmt.upper()}: Failed (see errors above)")
        
        logger.info("")
        logger.info(f"Exported {success_count}/{len(results)} formats successfully")
        logger.info(f"Output directory: {exporter.output_dir}")
        
        return 0 if success_count > 0 else 1
    
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1


if __name__ == '__main__':
    import sys
    sys.exit(main())
